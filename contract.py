import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
GENERATED_CONTRACTS_DIR = os.path.join(BASE_DIR, "generated_contracts")

os.makedirs(GENERATED_CONTRACTS_DIR, exist_ok=True)

# Number to words conversion
NUMBER_TO_WORDS = {
    1: "one (1)",
    2: "one (1) / two (2)",
    3: "one (1) / two (2) / three (3)",
    4: "one (1) / two (2) / three (3) / four (4)",
    5: "one (1) / two (2) / three (3) / four (4) / five (5)",
    6: "one (1) / two (2) / three (3) / four (4) / five (5) / six (6)",
    7: "one (1) / two (2) / three (3) / four (4) / five (5) / six (6) / seven (7)",
    8: "one (1) / two (2) / three (3) / four (4) / five (5) / six (6) / seven (7) / eight (8)",
    9: "one (1) / two (2) / three (3) / four (4) / five (5) / six (6) / seven (7) / eight (8) / nine (9)",
    10: "one (1) / two (2) / three (3) / four (4) / five (5) / six (6) / seven (7) / eight (8) / nine (9) / ten (10)"
}

def convert_number_to_words(num):
    """Convert number to words format"""
    try:
        number = int(num)
        return NUMBER_TO_WORDS.get(number, f"{number}")
    except (ValueError, TypeError):
        return "one (1)"

def add_paragraph(doc, text, bold=False, space_before=0, space_after=0):
    """Helper to add a paragraph with formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    
    # Add spacing
    if space_before > 0:
        para.paragraph_format.space_before = Pt(space_before)
    if space_after > 0:
        para.paragraph_format.space_after = Pt(space_after)
    
    return para

def generate_regular_contract(doc, data, num_text):
    """Generate Regular template contract"""
    
    # Header info
    add_paragraph(doc, f"Artist/vendor name: {data['contractor_name']}")
    add_paragraph(doc, f"Name of person signing the contract (if not the artist/vendor): {data.get('signer_name', '')}")
    add_paragraph(doc, f"Relationship to artist/vendor: {data.get('relationship_to_vendor', '')}")
    add_paragraph(doc, f"Address: {data['address']}")
    add_paragraph(doc, f"Email address: {data['email']}")
    add_paragraph(doc, f"Vendor account: {data.get('vendor_account', '')}", space_after=12)
    
    # Summary
    add_paragraph(doc, "Summary:", bold=True, space_before=12)
    add_paragraph(doc, f"Vendor will create and provide to Adobe {num_text} video(s) with content to promote select Adobe products.", space_after=12)
    
    # Deliverables
    add_paragraph(doc, "Deliverables:", bold=True, space_before=12)
    add_paragraph(doc, f"Vendor will provide Adobe with {num_text} pre-recorded video(s) that will be between 30 seconds and one minute in length that highlight Adobe products. Specific details, including Adobe product(s), will be selected by Adobe in writing. For each video, Vendor will (1) orally disclose the relationship between Vendor and Adobe and (2) include a clearly visible written overlay disclosing the relationship. Unless otherwise specified by Adobe in writing, each video's aspect ratio will be 9:16.")
    
    add_paragraph(doc, "Vendor will post the video(s) on various social media channels owned and controlled by the Vendor, which the parties will agree to in writing. [The video(s) must be authenticated via the CreatorIQ website for analytic purposes, with a 30-day Ad code for all video created on applicable social media platforms provided to Adobe to track performance.]", space_after=12)
    
    add_paragraph(doc, "For clarity, Adobe shall have the right to like, favorite, share, repost, redistribute, syndicate, amplify (paid promotion or allow listing) or otherwise use all video described hereunder in any manner enabled by the applicable platform. Adobe can use the video and may redistribute to other Adobe owned accounts, channels, and/or platforms. Vendor will allow 1 round of edits per video.", space_after=12)
    
    add_paragraph(doc, "In the event that all pre-recorded video(s) are not delivered, Adobe will pay a pro-rated rate for content delivered in accordance with this Agreement.", space_after=12)
    
    add_paragraph(doc, 'Beginning on the Effective Date, and concluding thirty (30) days after Vendor\'s publication of the video(s) with Adobe\'s authorization, Vendor will not provide services on behalf of, appear or participate in any advertising, publicity or promotion of, endorse, or authorize or permit the use of Vendor\'s Likeness in connection with the following (the "Restricted Category"): (a) any software and online creative development and cloud service companies (for clarity, the Restricted Category includes, without limitation, Spline, Womp, Canva, Affinity, CapCut, Autodesk, DaVinci, Final Cut Pro, Figma, Procreate, Capture One Pro); or (b) any product or service that in its advertising or publicity denigrates Adobe or its products. For clarity, the aforementioned does not preclude Vendor from merely appearing in any entertainment portion of any news, TV, or film program or attending an event, regardless of sponsorship.', space_after=12)
    
    # Delivery Schedule
    add_paragraph(doc, "Delivery Schedule:", bold=True, space_before=12)
    add_paragraph(doc, f"Unless otherwise directed in writing by Adobe, {data.get('due_date', '')}", space_after=12)
    
    # Price and End Date
    add_paragraph(doc, f"Price and currency: ${data.get('amount', '')} USD")
    add_paragraph(doc, f"End Date: {data.get('end_date', '')}")

def generate_campfire_contract(doc, data, num_text):
    """Generate Campfire template contract"""
    
    # Header info
    add_paragraph(doc, f"Artist/vendor name: {data['contractor_name']}")
    add_paragraph(doc, f"Name of person signing the contract (if not the artist/vendor): {data.get('signer_name', '')}")
    add_paragraph(doc, f"Relationship to artist/vendor: {data.get('relationship_to_vendor', '')}")
    add_paragraph(doc, f"Address: {data['address']}")
    add_paragraph(doc, f"Email address: {data['email']}")
    add_paragraph(doc, f"Vendor account: {data.get('vendor_account', '')}", space_after=12)
    
    # Summary
    add_paragraph(doc, "Summary:", bold=True, space_before=12)
    add_paragraph(doc, f"Vendor will create and provide to Adobe {num_text} video(s) with content to promote select Adobe products.", space_after=12)
    
    # Deliverables
    add_paragraph(doc, "Deliverables:", bold=True, space_before=12)
    add_paragraph(doc, f"Vendor will provide Adobe with {num_text} pre-recorded video(s) that will be between 30 seconds and one minute in length that highlight Adobe products. Specific details, including Adobe product(s), will be selected by Adobe in writing. For each video, Vendor will (1) orally disclose the relationship between Vendor and Adobe and (2) include a clearly visible written overlay disclosing the relationship. Unless otherwise specified by Adobe in writing, each video's aspect ratio will be 9:16.")
    
    add_paragraph(doc, "Vendor will post the video(s) on various social media channels owned and controlled by the Vendor, which the parties will agree to in writing. The video(s) must include a 30-day Ad code for all video created on applicable social media platforms provided to Adobe.", space_after=12)
    
    # Delivery Schedule
    add_paragraph(doc, "Delivery Schedule:", bold=True, space_before=12)
    add_paragraph(doc, f"Unless otherwise directed in writing by Adobe, {data.get('due_date', '')}", space_after=12)
    
    # Price and End Date
    add_paragraph(doc, f"Price and currency: ${data.get('amount', '')}")
    add_paragraph(doc, f"End Date: {data.get('end_date', '')}")

def generate_contract(fields: dict) -> str:
    """
    Generates contract Word document (.docx) using python-docx
    Required fields in dict:
    contractor_name, signer_name, relationship_to_vendor,
    address, email, vendor_account, service, amount, due_date, end_date,
    number_of_content, contract_type, po
    """

    # Create document
    doc = Document()
    
    # Set up page margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Convert number to words
    num_text = convert_number_to_words(fields.get('number_of_content', 1))
    
    # Generate appropriate template
    contract_type = fields.get('contract_type', 'regular').lower()
    if contract_type == 'campfire':
        generate_campfire_contract(doc, fields, num_text)
    else:
        generate_regular_contract(doc, fields, num_text)
    
    # Save document
    safe_name = fields["contractor_name"].replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"contract_{safe_name}_{timestamp}.docx"
    output_path = os.path.join(GENERATED_CONTRACTS_DIR, filename)
    
    doc.save(output_path)
    print(f"✅ Contract generated: {output_path}")
    
    return output_path


# Example local test
if __name__ == "__main__":
    data = {
        "contractor_name": "Bella Kotak",
        "signer_name": "",
        "relationship_to_vendor": "Self",
        "address": "123 Main St, San Francisco, CA 94102",
        "email": "bella@example.com",
        "vendor_account": "Needed",
        "service": "Video Production",
        "amount": "5000",
        "due_date": "December 31, 2025",
        "end_date": "December 31, 2025",
        "number_of_content": "3",
        "contract_type": "regular",  # or "campfire"
        "po": ""
    }

    path = generate_contract(data)
    print(f"✅ Contract generated at: {path}")
